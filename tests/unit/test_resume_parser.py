"""T11: 简历解析器测试 — parse_resume_document。

覆盖：PDF 解析、无效文件、PII 脱敏、空文件、非 PDF 文件。
"""

import pytest


def _get_parse_resume_document():
    from src.tools.resume_parser import parse_resume_document

    return parse_resume_document


class TestParseResumePDF:
    """PDF 简历解析。"""

    def test_parse_mock_resume(self, mock_resume_path):
        parse_resume_document = _get_parse_resume_document()
        result = parse_resume_document(mock_resume_path)
        assert isinstance(result, dict)
        # 应返回结构化数据（至少含 name 或 skills 字段）
        assert "name" in result or "skills" in result

    def test_parsed_result_no_raw_pii(self, mock_resume_path):
        """返回结果不应包含原始 PII 信息。"""
        parse_resume_document = _get_parse_resume_document()
        result = parse_resume_document(mock_resume_path)
        raw_text = str(result)
        # 手机号应被脱敏
        assert "13800138000" not in raw_text


class TestParseResumeInvalid:
    """无效文件处理。"""

    def test_nonexistent_file(self):
        parse_resume_document = _get_parse_resume_document()
        with pytest.raises(FileNotFoundError):
            parse_resume_document("/nonexistent/path/resume.pdf")

    def test_empty_file(self, tmp_path):
        empty_file = tmp_path / "empty.pdf"
        empty_file.write_bytes(b"")
        parse_resume_document = _get_parse_resume_document()
        with pytest.raises(ValueError):
            parse_resume_document(str(empty_file))


class TestParseResumeWord:
    """Word 简历解析（.docx）。"""

    def test_docx_file(self, tmp_path):
        parse_resume_document = _get_parse_resume_document()
        # 创建最小 docx 文件
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("张三 - Python 开发工程师")
            doc.add_paragraph("技能：Python, FastAPI, PostgreSQL")
            doc.save(str(tmp_path / "test.docx"))
            result = parse_resume_document(str(tmp_path / "test.docx"))
            assert isinstance(result, dict)
        except ImportError:
            pytest.skip("python-docx not installed")
